import os
from typing import Dict, Any, Optional
from pathlib import Path

class ExcelNativeRenderer:
    """Renders Excel templates with native openpyxl functionality."""

    def render(self, template_data: Dict, params: Dict) -> bytes:
        """Render Excel template with parameters."""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
        except ImportError:
            raise ImportError("openpyxl is required for Excel template rendering. Install with: pip install openpyxl")

        wb = Workbook()

        # Remove default sheet
        wb.remove(wb.active)

        sheets_config = template_data.get('sheets', {})

        for sheet_name, sheet_config in sheets_config.items():
            ws = wb.create_sheet(sheet_name)

            # Render cells
            cells_config = sheet_config.get('cells', {})
            for cell_ref, cell_config in cells_config.items():
                self._render_cell(ws, cell_ref, cell_config, params)

            # Render tables
            tables_config = sheet_config.get('tables', {})
            for table_name, table_config in tables_config.items():
                self._create_table(ws, table_config)

            # Render charts
            charts_config = sheet_config.get('charts', {})
            for chart_name, chart_config in charts_config.items():
                self._create_chart(ws, chart_config)

        # Apply global styles
        global_styles = template_data.get('global_styles', {})
        self._apply_global_styles(wb, global_styles)

        # Save to bytes
        from io import BytesIO
        bio = BytesIO()
        wb.save(bio)
        bio.seek(0)
        return bio.getvalue()

    def _render_cell(self, ws, cell_ref: str, cell_config: Dict, params: Dict) -> None:
        """Render a single cell with value and styling."""
        cell = ws[cell_ref]

        # Render value with parameter substitution
        value = cell_config.get('value', '')
        if isinstance(value, str):
            for param_key, param_value in params.items():
                placeholder = f"{{{{{param_key}}}}}"
                value = value.replace(placeholder, str(param_value))

        # Handle formulas
        if 'formula' in cell_config:
            formula = cell_config['formula']
            # Substitute parameters in formulas too
            for param_key, param_value in params.items():
                placeholder = f"{{{{{param_key}}}}}"
                formula = formula.replace(placeholder, str(param_value))
            cell.value = formula
        else:
            cell.value = value

        # Apply styles
        if 'style' in cell_config:
            self._apply_cell_style(cell, cell_config['style'])

    def _apply_cell_style(self, cell, style_config: Dict) -> None:
        """Apply styling to a cell."""
        try:
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        except ImportError:
            return

        # Font styling
        font_kwargs = {}
        if 'bold' in style_config:
            font_kwargs['bold'] = style_config['bold']
        if 'italic' in style_config:
            font_kwargs['italic'] = style_config['italic']
        if 'font_size' in style_config:
            font_kwargs['size'] = style_config['font_size']
        if 'font_color' in style_config:
            font_kwargs['color'] = style_config['font_color']

        if font_kwargs:
            cell.font = Font(**font_kwargs)

        # Fill/background
        if 'fill' in style_config:
            fill_config = style_config['fill']
            if fill_config.get('pattern') == 'solid' and 'color' in fill_config:
                cell.fill = PatternFill(start_color=fill_config['color'],
                                      end_color=fill_config['color'],
                                      fill_type='solid')

        # Alignment
        if 'alignment' in style_config:
            align_config = style_config['alignment']
            cell.alignment = Alignment(horizontal=align_config)

        # Border
        if 'border' in style_config:
            border_side = Side(style=style_config['border'])
            cell.border = Border(left=border_side, right=border_side,
                               top=border_side, bottom=border_side)

    def _create_table(self, ws, table_config: Dict) -> None:
        """Create an Excel table."""
        try:
            from openpyxl.worksheet.table import Table, TableStyleInfo
        except ImportError:
            return

        table_range = table_config.get('range', 'A1:B2')
        table = Table(displayName="Table1", ref=table_range)

        # Apply table style
        style_config = table_config.get('style', {})
        if 'header_fill' in style_config:
            table.tableStyleInfo = TableStyleInfo(
                name="TableStyleMedium9",
                showFirstColumn=False,
                showLastColumn=False,
                showRowStripes=True,
                showColumnStripes=True
            )

        ws.add_table(table)

    def _create_chart(self, ws, chart_config: Dict) -> None:
        """Create a chart (placeholder for future implementation)."""
        # Chart creation would require more complex implementation
        # For now, just skip
        pass

    def _apply_global_styles(self, wb, global_styles: Dict) -> None:
        """Apply global styles to the workbook."""
        # Implementation for global styles
        pass


class WordNativeRenderer:
    """Renders Word templates with native python-docx functionality."""

    def render(self, template_data: Dict, params: Dict) -> bytes:
        """Render Word template with parameters."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            raise ImportError("python-docx is required for Word template rendering. Install with: pip install python-docx")

        doc = Document()

        # Apply document settings
        doc_config = template_data.get('document', {})
        settings = doc_config.get('settings', {})
        self._apply_document_settings(doc, settings)

        # Apply custom styles
        styles_config = doc_config.get('styles', {})
        self._apply_styles(doc, styles_config)

        # Render sections
        sections = doc_config.get('sections', [])
        for section in sections:
            self._render_section(doc, section, params)

        # Apply footer if configured
        footer_config = doc_config.get('footer', {})
        if footer_config:
            self._apply_footer(doc, footer_config, params)

        # Save to bytes
        from io import BytesIO
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio.getvalue()

    def _apply_document_settings(self, doc, settings: Dict) -> None:
        """Apply document-level settings."""
        try:
            from docx.shared import Inches
            from docx.enum.section import WD_ORIENT

            sections = doc.sections
            for section in sections:
                if 'page_size' in settings:
                    # A4 is default, could be extended for other sizes
                    pass

                if 'orientation' in settings:
                    if settings['orientation'] == 'landscape':
                        section.orientation = WD_ORIENT.LANDSCAPE

                margins = settings.get('margins', {})
                if 'top' in margins:
                    section.top_margin = Inches(margins['top'])
                if 'bottom' in margins:
                    section.bottom_margin = Inches(margins['bottom'])
                if 'left' in margins:
                    section.left_margin = Inches(margins['left'])
                if 'right' in margins:
                    section.right_margin = Inches(margins['right'])
        except ImportError:
            pass

    def _apply_styles(self, doc, styles_config: Dict) -> None:
        """Apply custom styles to the document."""
        try:
            from docx.enum.style import WD_STYLE_TYPE
            from docx.shared import Pt, RGBColor

            for style_name, style_config in styles_config.items():
                style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

                # Font properties
                font = style.font
                if 'font_size' in style_config:
                    font.size = Pt(style_config['font_size'])
                if 'bold' in style_config and style_config['bold']:
                    font.bold = True
                if 'italic' in style_config and style_config['italic']:
                    font.italic = True
                if 'underline' in style_config and style_config['underline']:
                    font.underline = True
                if 'font_color' in style_config:
                    # Convert hex to RGB if needed
                    font.color.rgb = self._hex_to_rgb(style_config['font_color'])

                # Paragraph properties
                para_format = style.paragraph_format
                if 'alignment' in style_config:
                    align_map = {
                        'left': WD_ALIGN_PARAGRAPH.LEFT,
                        'center': WD_ALIGN_PARAGRAPH.CENTER,
                        'right': WD_ALIGN_PARAGRAPH.RIGHT,
                        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
                    }
                    para_format.alignment = align_map.get(style_config['alignment'], WD_ALIGN_PARAGRAPH.LEFT)

                spacing = style_config.get('spacing', {})
                if 'before' in spacing:
                    para_format.space_before = Pt(spacing['before'])
                if 'after' in spacing:
                    para_format.space_after = Pt(spacing['after'])
        except ImportError:
            pass

    def _render_section(self, doc, section: Dict, params: Dict) -> None:
        """Render a document section."""
        section_type = section.get('type', 'paragraph')

        if section_type == 'header':
            level = section.get('level', 1)
            content = self._render_content(section.get('content', ''), params)
            doc.add_heading(content, level=level)

        elif section_type == 'paragraph':
            content = self._render_content(section.get('content', ''), params)
            style_name = section.get('style')
            p = doc.add_paragraph(content, style=style_name)

        elif section_type == 'table':
            self._render_table(doc, section, params)

        elif section_type == 'list':
            self._render_list(doc, section, params)

        elif section_type == 'section_break':
            # Add page break
            doc.add_page_break()

    def _render_table(self, doc, table_config: Dict, params: Dict) -> None:
        """Render a table in the document."""
        title = table_config.get('title', '')
        if title:
            doc.add_heading(self._render_content(title, params), level=3)

        headers = table_config.get('headers', [])
        rows_data = table_config.get('rows', [])

        if headers:
            table = doc.add_table(rows=1, cols=len(headers))
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = self._render_content(header, params)

            # Add data rows
            for row_data in rows_data:
                row_cells = table.add_row().cells
                for i, cell_data in enumerate(row_data):
                    row_cells[i].text = self._render_content(str(cell_data), params)

    def _render_list(self, doc, list_config: Dict, params: Dict) -> None:
        """Render a list in the document."""
        list_type = list_config.get('list_type', 'bullet')
        items = list_config.get('items', [])

        for item in items:
            content = self._render_content(item, params)
            doc.add_paragraph(content, style='List Bullet' if list_type == 'bullet' else 'List Number')

    def _apply_footer(self, doc, footer_config: Dict, params: Dict) -> None:
        """Apply footer to all sections."""
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            for section in doc.sections:
                footer = section.footer
                footer_para = footer.paragraphs[0]
                footer_para.text = self._render_content(footer_config.get('content', ''), params)

                # Apply footer style
                style_config = footer_config.get('style', {})
                if 'font_size' in style_config:
                    footer_para.style.font.size = Pt(style_config['font_size'])
                if 'italic' in style_config and style_config['italic']:
                    footer_para.style.font.italic = True
                if 'alignment' in style_config:
                    align_map = {
                        'left': WD_ALIGN_PARAGRAPH.LEFT,
                        'center': WD_ALIGN_PARAGRAPH.CENTER,
                        'right': WD_ALIGN_PARAGRAPH.RIGHT
                    }
                    footer_para.alignment = align_map.get(style_config['alignment'], WD_ALIGN_PARAGRAPH.CENTER)
        except ImportError:
            pass

    def _render_content(self, content: str, params: Dict) -> str:
        """Render content with parameter substitution."""
        if not isinstance(content, str):
            return str(content)

        result = content
        for param_key, param_value in params.items():
            placeholder = f"{{{{{param_key}}}}}"
            result = result.replace(placeholder, str(param_value))

        return result

    def _hex_to_rgb(self, hex_color: str) -> 'RGBColor':
        """Convert hex color to RGBColor."""
        try:
            from docx.shared import RGBColor
            hex_color = hex_color.lstrip('#')
            return RGBColor(int(hex_color[0:2], 16),
                          int(hex_color[2:4], 16),
                          int(hex_color[4:6], 16))
        except (ImportError, ValueError):
            return RGBColor(0, 0, 0)  # Default to black


class Jinja2Renderer:
    """Renders templates using Jinja2 template engine."""

    def render(self, template_data: Dict, params: Dict) -> str:
        """Render Jinja2 template with parameters."""
        try:
            from jinja2 import Template
        except ImportError:
            raise ImportError("jinja2 is required for Jinja2 template rendering. Install with: pip install jinja2")

        template_content = template_data.get('content', '')
        template = Template(template_content)
        return template.render(**params)