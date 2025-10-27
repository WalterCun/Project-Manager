import os
from abc import ABC, abstractmethod
from typing import Dict, Any, List
from .models import TemplateManager

class FileRenderer(ABC):
    """Abstract base class for file renderers."""

    @abstractmethod
    def render(self, content: str, output_path: str) -> None:
        """Render content to a file."""
        pass

    @abstractmethod
    def get_extension(self) -> str:
        """Return the file extension this renderer handles."""
        pass

class DocxRenderer(FileRenderer):
    """Renderer for .docx files using python-docx."""

    def get_extension(self) -> str:
        return 'docx'

    def render(self, content: str, output_path: str) -> None:
        from docx import Document
        from docx.shared import Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        import json

        doc = Document()

        # Try to parse content as structured JSON
        try:
            if '{' in content and '}' in content:
                start = content.find('{')
                end = content.rfind('}') + 1
                if start != -1 and end > start:
                    json_str = content[start:end]
                    data = json.loads(json_str)
                    self._render_structured_docx(doc, data)
                else:
                    self._render_simple_docx(doc, content)
            else:
                self._render_simple_docx(doc, content)
        except Exception as e:
            print(f"Warning: Failed to render structured DOCX: {e}")
            self._render_simple_docx(doc, content)

        doc.save(output_path)

    def _render_simple_docx(self, doc, content: str) -> None:
        """Render simple text content to DOCX."""
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                doc.add_paragraph(para_text.strip())

    def _render_structured_docx(self, doc, data: Dict[str, Any]) -> None:
        """Render structured JSON template to DOCX."""
        from docx.shared import Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT

        # Apply document settings
        if 'document' in data and 'settings' in data['document']:
            self._apply_document_settings(doc, data['document']['settings'])

        # Get global styles
        global_styles = {}
        if 'document' in data and 'styles' in data['document']:
            global_styles = data['document']['styles']

        # Render sections
        if 'document' in data and 'sections' in data['document']:
            self._render_sections(doc, data['document']['sections'], global_styles)

        # Add footer if specified
        if 'document' in data and 'footer' in data['document']:
            self._add_footer(doc, data['document']['footer'])

    def _apply_document_settings(self, doc, settings: Dict[str, Any]) -> None:
        """Apply document-level settings."""
        # Note: python-docx has limited support for document settings
        # These would need to be applied after document creation
        pass

    def _render_sections(self, doc, sections: List[Dict[str, Any]], global_styles: Dict[str, Any]) -> None:
        """Render document sections."""
        for section in sections:
            section_type = section.get('type', 'paragraph')

            if section_type == 'header':
                self._render_header(doc, section, global_styles)
            elif section_type == 'paragraph':
                self._render_paragraph(doc, section, global_styles)
            elif section_type == 'table':
                self._render_table(doc, section, global_styles)
            elif section_type == 'list':
                self._render_list(doc, section, global_styles)
            elif section_type == 'table_of_contents':
                self._render_table_of_contents(doc, section, global_styles)
            elif section_type == 'section_break':
                self._add_section_break(doc, section)

    def _render_header(self, doc, section: Dict[str, Any], global_styles: Dict[str, Any]) -> None:
        """Render a header section."""
        content = section.get('content', '')
        level = section.get('level', 1)
        style_name = section.get('style', 'heading1_style')

        # Get style configuration
        style = global_styles.get(style_name, {})

        # Create paragraph with content
        para = doc.add_paragraph(content)

        # Apply styling
        self._apply_paragraph_style(para, style)

        # Apply level-based formatting
        if level == 1:
            para.style = 'Heading 1'
        elif level == 2:
            para.style = 'Heading 2'
        elif level == 3:
            para.style = 'Heading 3'

    def _render_paragraph(self, doc, section: Dict[str, Any], global_styles: Dict[str, Any]) -> None:
        """Render a paragraph section."""
        content = section.get('content', '')
        style_name = section.get('style', 'normal_style')
        style = global_styles.get(style_name, {})

        para = doc.add_paragraph(content)
        self._apply_paragraph_style(para, style)

    def _render_table(self, doc, section: Dict[str, Any], global_styles: Dict[str, Any]) -> None:
        """Render a table section."""
        headers = section.get('headers', [])
        rows = section.get('rows', [])
        style = section.get('style', {})

        if not headers and not rows:
            return

        # Create table
        num_cols = max(len(headers), max(len(row) for row in rows) if rows else 0)
        table = doc.add_table(rows=len(rows) + (1 if headers else 0), cols=num_cols)

        # Add headers
        if headers:
            header_row = table.rows[0]
            for i, header in enumerate(headers):
                if i < len(header_row.cells):
                    header_row.cells[i].text = header
                    self._apply_cell_style(header_row.cells[i], style.get('header_style', {}))

        # Add data rows
        for row_idx, row in enumerate(rows):
            table_row = table.rows[row_idx + (1 if headers else 0)]
            for col_idx, cell_content in enumerate(row):
                if col_idx < len(table_row.cells):
                    table_row.cells[col_idx].text = cell_content
                    self._apply_cell_style(table_row.cells[col_idx], style.get('cell_style', {}))

        # Apply table-level styling
        if style.get('border'):
            # Apply borders to all cells
            for row in table.rows:
                for cell in row.cells:
                    self._apply_cell_borders(cell, style['border'])

    def _render_list(self, doc, section: Dict[str, Any], global_styles: Dict[str, Any]) -> None:
        """Render a list section."""
        items = section.get('items', [])
        list_type = section.get('list_type', 'bullet')
        style_name = section.get('style', 'normal_style')
        style = global_styles.get(style_name, {})

        for item in items:
            para = doc.add_paragraph(item, style='List Bullet' if list_type == 'bullet' else 'List Number')
            self._apply_paragraph_style(para, style)

    def _render_table_of_contents(self, doc, section: Dict[str, Any], global_styles: Dict[str, Any]) -> None:
        """Render table of contents (placeholder)."""
        title = section.get('title', 'Tabla de Contenido')
        style_name = section.get('style', 'heading1_style')
        style = global_styles.get(style_name, {})

        para = doc.add_paragraph(f"[{title} - Auto-generated]")
        self._apply_paragraph_style(para, style)

    def _add_section_break(self, doc, section: Dict[str, Any]) -> None:
        """Add a section break."""
        break_type = section.get('break_type', 'page')
        # Note: python-docx has limited section break support
        doc.add_paragraph(f"[Section Break: {break_type}]")

    def _add_footer(self, doc, footer_config: Dict[str, Any]) -> None:
        """Add document footer."""
        # Note: python-docx footer support is limited in simple usage
        pass

    def _apply_paragraph_style(self, paragraph, style: Dict[str, Any]) -> None:
        """Apply styling to a paragraph."""
        from docx.shared import Inches, Cm

        if 'font_size' in style:
            paragraph.style.font.size = Inches(style['font_size'] / 72)  # Convert points to inches

        if 'bold' in style and style['bold']:
            paragraph.style.font.bold = True

        if 'italic' in style and style['italic']:
            paragraph.style.font.italic = True

        if 'underline' in style and style['underline']:
            paragraph.style.font.underline = True

        if 'font_color' in style:
            # Note: Font color requires RGB values
            pass

        if 'alignment' in style:
            align_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY
            }
            if style['alignment'] in align_map:
                paragraph.alignment = align_map[style['alignment']]

        if 'spacing' in style:
            spacing = style['spacing']
            if 'before' in spacing:
                paragraph.paragraph_format.space_before = Inches(spacing['before'] / 72)
            if 'after' in spacing:
                paragraph.paragraph_format.space_after = Inches(spacing['after'] / 72)

        if 'line_spacing' in style:
            # Line spacing is more complex in python-docx
            pass

    def _apply_cell_style(self, cell, style: Dict[str, Any]) -> None:
        """Apply styling to a table cell."""
        for paragraph in cell.paragraphs:
            self._apply_paragraph_style(paragraph, style)

    def _apply_cell_borders(self, cell, border_type: str) -> None:
        """Apply borders to a table cell."""
        # Note: Border styling in python-docx is complex
        # This is a placeholder for future implementation
        pass

class XlsxRenderer(FileRenderer):
    """Renderer for .xlsx files using openpyxl."""

    def get_extension(self) -> str:
        return 'xlsx'

    def render(self, content: str, output_path: str) -> None:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill
        import json

        wb = Workbook()

        # Try to parse content as structured data
        try:
            # If content contains JSON-like structure, parse it
            if '{' in content and '}' in content:
                # Simple JSON detection
                start = content.find('{')
                end = content.rfind('}') + 1
                if start != -1 and end > start:
                    json_str = content[start:end]
                    data = json.loads(json_str)
                    self._render_excel_template(wb, data)
                else:
                    self._render_simple_content(wb.active, content)
            else:
                self._render_simple_content(wb.active, content)
        except Exception as e:
            print(f"Warning: Failed to render Excel template: {e}")
            self._render_simple_content(wb.active, content)

        wb.save(output_path)

    def _render_simple_content(self, ws, content: str) -> None:
        """Render simple text content."""
        from openpyxl.styles import Font, Alignment

        # Split content into lines
        lines = content.split('\n')

        # Title
        if lines:
            ws['A1'] = lines[0]
            ws['A1'].font = Font(size=16, bold=True)
            ws['A1'].alignment = Alignment(horizontal='center')

        # Content
        for i, line in enumerate(lines[1:], 2):
            if line.strip():
                ws[f'A{i}'] = line
                if line.startswith('#'):
                    ws[f'A{i}'].font = Font(bold=True)
                elif line.startswith('##'):
                    ws[f'A{i}'].font = Font(size=14, bold=True)

        # Auto-adjust column width
        ws.column_dimensions['A'].width = 50

    def _render_excel_template(self, wb, data: Dict[str, Any]) -> None:
        """Render a multi-sheet Excel template."""
        from openpyxl.styles import Font, Alignment, PatternFill

        # Handle multi-sheet templates
        if 'sheets' in data:
            self._render_multi_sheet_excel(wb, data)
        else:
            # Fallback to single sheet
            self._render_structured_data(wb.active, data)

    def _render_multi_sheet_excel(self, wb, data: Dict[str, Any]) -> None:
        """Render multi-sheet Excel template."""
        from openpyxl.styles import Font, Alignment, PatternFill

        sheets = data.get('sheets', {})
        global_styles = data.get('global_styles', {})

        # Remove default sheet if multiple sheets are defined
        if len(sheets) > 1:
            wb.remove(wb.active)

        for sheet_name, sheet_config in sheets.items():
            ws = wb.create_sheet(title=sheet_name)

            # Apply sheet description as comment if available
            if 'description' in sheet_config:
                ws.cell(1, 1).comment = sheet_config['description']

            # Render cells
            if 'cells' in sheet_config:
                self._render_cells(ws, sheet_config['cells'], global_styles)

            # Render tables
            if 'tables' in sheet_config:
                self._render_tables(ws, sheet_config['tables'], global_styles)

            # Render charts (placeholder for future implementation)
            if 'charts' in sheet_config:
                self._render_charts_placeholder(ws, sheet_config['charts'])

    def _render_cells(self, ws, cells: Dict[str, Any], global_styles: Dict[str, Any]) -> None:
        """Render individual cells with styles."""
        from openpyxl.styles import Font, Alignment, PatternFill

        for cell_ref, cell_config in cells.items():
            if 'value' in cell_config:
                ws[cell_ref] = cell_config['value']
            elif 'formula' in cell_config:
                ws[cell_ref] = cell_config['formula']

            # Apply styles
            if 'style' in cell_config:
                style = cell_config['style']
                cell = ws[cell_ref]

                if 'bold' in style and style['bold']:
                    cell.font = Font(bold=True)
                if 'font_size' in style:
                    if not cell.font:
                        cell.font = Font()
                    cell.font = Font(size=style['font_size'], bold=cell.font.bold if cell.font else False)
                if 'font_color' in style:
                    if not cell.font:
                        cell.font = Font()
                    cell.font = Font(color=style['font_color'], size=cell.font.size if cell.font else 11)
                if 'fill' in style:
                    fill_config = style['fill']
                    if fill_config.get('pattern') == 'solid' and 'color' in fill_config:
                        cell.fill = PatternFill(start_color=fill_config['color'], end_color=fill_config['color'], fill_type="solid")
                if 'alignment' in style:
                    align_config = style['alignment']
                    cell.alignment = Alignment(horizontal=align_config.get('horizontal', 'left'))

    def _render_tables(self, ws, tables: Dict[str, Any], global_styles: Dict[str, Any]) -> None:
        """Render table structures."""
        from openpyxl.styles import Font, PatternFill

        for table_name, table_config in tables.items():
            range_str = table_config.get('range', '')
            if not range_str:
                continue

            # Parse range (e.g., "B7:D9")
            try:
                start_col, start_row = self._parse_cell_ref(range_str.split(':')[0])
                end_col, end_row = self._parse_cell_ref(range_str.split(':')[1])

                # Apply table styling
                style = table_config.get('style', {})

                # Style headers
                if table_config.get('headers', False):
                    for col in range(start_col, end_col + 1):
                        cell = ws.cell(start_row, col)
                        if style.get('header_fill'):
                            cell.fill = PatternFill(start_color=style['header_fill'], end_color=style['header_fill'], fill_type="solid")
                        if style.get('header_font_color'):
                            cell.font = Font(color=style['header_font_color'], bold=True)

                # Apply borders if specified
                if style.get('border'):
                    for row in range(start_row, end_row + 1):
                        for col in range(start_col, end_col + 1):
                            cell = ws.cell(row, col)
                            # Add border styling here when needed

            except Exception as e:
                print(f"Warning: Failed to render table {table_name}: {e}")

    def _render_charts_placeholder(self, ws, charts: Dict[str, Any]) -> None:
        """Add placeholder text for charts (future implementation)."""
        for chart_name, chart_config in charts.items():
            position = chart_config.get('position', 'A1')
            ws[position] = f"[Chart: {chart_name}] - {chart_config.get('type', 'unknown')}"

    def _parse_cell_ref(self, cell_ref: str) -> tuple:
        """Parse cell reference like 'B7' to (col, row)."""
        import re
        match = re.match(r'([A-Z]+)(\d+)', cell_ref)
        if match:
            col_str, row_str = match.groups()
            col = 0
            for i, char in enumerate(col_str):
                col += (ord(char) - ord('A') + 1) * (26 ** (len(col_str) - i - 1))
            return (col, int(row_str))
        return (1, 1)

    def _render_structured_data(self, ws, data: Dict[str, Any]) -> None:
        """Render structured data with headers and tables (fallback method)."""
        from openpyxl.styles import Font, Alignment, PatternFill

        row = 1

        # Title
        if 'titulo' in data:
            ws[f'A{row}'] = data['titulo']
            ws[f'A{row}'].font = Font(size=16, bold=True)
            ws[f'A{row}'].alignment = Alignment(horizontal='center')
            ws.merge_cells(f'A{row}:E{row}')
            row += 2

        # Company and date
        if 'empresa' in data:
            ws[f'A{row}'] = f"Empresa: {data['empresa']}"
            ws[f'A{row}'].font = Font(italic=True)
            row += 1

        if 'fecha' in data and 'hora' in data:
            ws[f'A{row}'] = f"Fecha: {data['fecha']} | Hora: {data['hora']}"
            ws[f'A{row}'].font = Font(italic=True)
            row += 2

        # Main content sections
        sections = ['introduccion', 'objetivo', 'contenido', 'conclusiones']
        for section in sections:
            if section in data:
                ws[f'A{row}'] = section.replace('_', ' ').title()
                ws[f'A{row}'].font = Font(size=14, bold=True)
                ws[f'A{row}'].fill = PatternFill(start_color="E6E6FA", end_color="E6E6FA", fill_type="solid")
                row += 1

                content_value = data[section]
                if isinstance(content_value, str):
                    content_lines = content_value.split('\n')
                    for line in content_lines:
                        if line.strip():
                            ws[f'A{row}'] = line
                            row += 1
                elif isinstance(content_value, dict) and content_value:
                    # Render as table
                    col = 1
                    for header in content_value.keys():
                        ws.cell(row=row, column=col, value=header)
                        ws.cell(row=row, column=col).font = Font(bold=True)
                        ws.cell(row=row, column=col).fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
                        col += 1
                    row += 1

                    max_rows = max(len(v) if isinstance(v, list) else 1 for v in content_value.values())
                    for i in range(max_rows):
                        col = 1
                        for header, values in content_value.items():
                            if isinstance(values, list) and i < len(values):
                                ws.cell(row=row, column=col, value=values[i])
                            elif not isinstance(values, list):
                                ws.cell(row=row, column=col, value=values)
                            col += 1
                        row += 1
                row += 1

        # Footer
        ws[f'A{row}'] = "Generado automáticamente por Project Manager"
        ws[f'A{row}'].font = Font(size=10, italic=True)
        ws[f'A{row}'].alignment = Alignment(horizontal='center')
        ws.merge_cells(f'A{row}:E{row}')

        # Auto-adjust column widths
        for col in ['A', 'B', 'C', 'D', 'E']:
            ws.column_dimensions[col].width = 20

class HtmlRenderer(FileRenderer):
    """Renderer for .html files."""

    def get_extension(self) -> str:
        return 'html'

    def render(self, content: str, output_path: str) -> None:
        import json

        # Try to parse content as structured JSON
        try:
            if '{' in content and '}' in content:
                start = content.find('{')
                end = content.rfind('}') + 1
                if start != -1 and end > start:
                    json_str = content[start:end]
                    data = json.loads(json_str)
                    self._render_structured_html(data, output_path)
                    return
        except Exception:
            pass

        # Fallback to simple HTML rendering
        self._render_simple_html(content, output_path)

    def _render_simple_html(self, content: str, output_path: str) -> None:
        """Render simple text content to HTML."""
        html_content = f"""
<!DOCTYPE html>
<html>
<head>
    <title>Generated Document</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 20px; }}
        p {{ line-height: 1.6; }}
    </style>
</head>
<body>
    <p>{content}</p>
</body>
</html>
        """
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

    def _render_structured_html(self, data: Dict[str, Any], output_path: str) -> None:
        """Render structured JSON template to HTML."""
        # Handle advanced HTML template format
        if 'document' in data:
            self._render_advanced_html(data['document'], output_path)
        else:
            self._render_simple_html(str(data), output_path)

    def _render_advanced_html(self, document: Dict[str, Any], output_path: str) -> None:
        """Render advanced HTML document with sections."""
        html_parts = []

        # HTML header
        html_parts.append("<!DOCTYPE html>")
        html_parts.append('<html lang="es">')
        html_parts.append("<head>")
        html_parts.append(f'    <title>{document.get("title", "Generated Document")}</title>')
        html_parts.append('    <meta charset="UTF-8">')
        html_parts.append('    <meta name="viewport" content="width=device-width, initial-scale=1.0">')

        # Include CSS frameworks
        includes = document.get('includes', {})
        for lib_name, lib_url in includes.items():
            if 'bootstrap' in lib_name.lower():
                html_parts.append(f'    <link href="{lib_url}" rel="stylesheet">')

        # Custom styles
        html_parts.append('    <style>')
        styles = document.get('styles', {})
        for selector, style_dict in styles.items():
            html_parts.append(f'        {selector} {{')
            for prop, value in style_dict.items():
                html_parts.append(f'            {prop.replace("_", "-")}: {value};')
            html_parts.append('        }')
        html_parts.append('    </style>')

        html_parts.append("</head>")
        html_parts.append("<body>")

        # Render sections
        sections = document.get('sections', [])
        for section in sections:
            section_type = section.get('type', 'paragraph')

            if section_type == 'header':
                content = section.get('content', '')
                subtitle = section.get('subtitle', '')
                metadata = section.get('metadata', '')

                html_parts.append('    <div class="header">')
                html_parts.append(f'        <h1>{content}</h1>')
                if subtitle:
                    html_parts.append(f'        <h2>{subtitle}</h2>')
                if metadata:
                    html_parts.append(f'        <p>{metadata}</p>')
                html_parts.append('    </div>')

            elif section_type == 'metric_grid':
                columns = section.get('columns', 3)
                metrics = section.get('metrics', [])

                html_parts.append(f'    <div class="row">')
                for metric in metrics:
                    html_parts.append('        <div class="col-md-{12//columns}">')
                    html_parts.append('            <div class="metric-card text-center">')
                    html_parts.append(f'                <i class="{metric.get("icon", "fas fa-chart-bar")}"></i>')
                    html_parts.append(f'                <h5>{metric.get("title", "")}</h5>')
                    html_parts.append(f'                <h3>{metric.get("value", "")}</h3>')
                    html_parts.append(f'                <span class="status-badge {metric.get("status_class", "")}">{metric.get("status", "")}</span>')
                    html_parts.append(f'                <p>{metric.get("description", "")}</p>')
                    html_parts.append('            </div>')
                    html_parts.append('        </div>')
                html_parts.append('    </div>')

            elif section_type == 'content_section':
                title = section.get('title', '')
                content = section.get('content', '')
                icon = section.get('icon', '')

                html_parts.append('    <div class="metric-card">')
                html_parts.append(f'        <h4><i class="{icon}"></i> {title}</h4>')
                html_parts.append(f'        <p>{content}</p>')
                html_parts.append('    </div>')

            elif section_type == 'chart_section':
                title = section.get('title', '')
                chart_type = section.get('chart_type', 'bar')
                chart_id = section.get('chart_id', 'chart1')

                html_parts.append('    <div class="metric-card">')
                html_parts.append(f'        <h4>{title}</h4>')
                html_parts.append(f'        <div class="chart-container">')
                html_parts.append(f'            <canvas id="{chart_id}"></canvas>')
                html_parts.append('        </div>')
                html_parts.append('    </div>')

            elif section_type == 'analysis_section':
                title = section.get('title', '')
                content = section.get('content', '')
                highlights = section.get('highlights', [])
                icon = section.get('icon', '')

                html_parts.append('    <div class="metric-card">')
                html_parts.append(f'        <h4><i class="{icon}"></i> {title}</h4>')
                html_parts.append(f'        <p>{content}</p>')
                if highlights:
                    html_parts.append('        <ul>')
                    for highlight in highlights:
                        html_parts.append(f'            <li>{highlight}</li>')
                    html_parts.append('        </ul>')
                html_parts.append('    </div>')

            elif section_type == 'conclusions_section':
                title = section.get('title', '')
                content = section.get('content', '')
                recommendations = section.get('recommendations', [])
                icon = section.get('icon', '')

                html_parts.append('    <div class="metric-card">')
                html_parts.append(f'        <h4><i class="{icon}"></i> {title}</h4>')
                html_parts.append(f'        <p>{content}</p>')
                if recommendations:
                    html_parts.append('        <ul>')
                    for rec in recommendations:
                        html_parts.append(f'            <li>{rec}</li>')
                    html_parts.append('        </ul>')
                html_parts.append('    </div>')

            elif section_type == 'footer':
                content = section.get('content', '')
                style = section.get('style', {})

                style_str = ' '.join([f'{k.replace("_", "-")}: {v};' for k, v in style.items()])
                html_parts.append(f'    <footer style="{style_str}">')
                html_parts.append(f'        <p>{content}</p>')
                html_parts.append('    </footer>')

        # Add scripts
        scripts = document.get('scripts', {})
        if scripts:
            html_parts.append('    <script src="https://cdn.jsdelivr.net/npm/chart.js"></script>')
            html_parts.append('    <script>')
            html_parts.append('        function initializeCharts() {')
            html_parts.append('            // Chart initialization code would go here')
            html_parts.append('        }')
            html_parts.append('        document.addEventListener("DOMContentLoaded", initializeCharts);')
            html_parts.append('    </script>')

        html_parts.append("</body>")
        html_parts.append("</html>")

        # Write HTML file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(html_parts))

class MdRenderer(FileRenderer):
    """Renderer for .md files."""

    def get_extension(self) -> str:
        return 'md'

    def render(self, content: str, output_path: str) -> None:
        with open(output_path, 'w') as f:
            f.write(content)

class TxtRenderer(FileRenderer):
    """Renderer for .txt files."""

    def get_extension(self) -> str:
        return 'txt'

    def render(self, content: str, output_path: str) -> None:
        with open(output_path, 'w') as f:
            f.write(content)

class RendererFactory:
    """Factory to get the appropriate renderer."""

    _renderers = {
        'docx': DocxRenderer(),
        'xlsx': XlsxRenderer(),
        'html': HtmlRenderer(),
        'md': MdRenderer(),
        'txt': TxtRenderer(),
    }

    @classmethod
    def get_renderer(cls, extension: str) -> FileRenderer:
        """Get renderer for the given extension."""
        if extension not in cls._renderers:
            raise ValueError(f"Unsupported extension: {extension}")
        return cls._renderers[extension]

    @classmethod
    def add_renderer(cls, extension: str, renderer: FileRenderer) -> None:
        """Add a new renderer for a custom extension."""
        cls._renderers[extension] = renderer